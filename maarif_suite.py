import streamlit as st
import google.generativeai as genai
from groq import Groq
import tempfile
import os
from io import BytesIO 
from docx import Document 

# --- 1. GÜVENLİK VE API AYARLARI ---

GOOGLE_API_KEY = st.secrets.get("GOOGLE_API_KEY")
GROQ_API_KEY = st.secrets.get("GROQ_API_KEY")

if not GOOGLE_API_KEY or not GROQ_API_KEY:
    st.error("HATA: Google API Anahtarı ve/veya Groq API Anahtarı bulunamadı! Lütfen secrets dosyasını kontrol edin.")
    st.stop()

try:
    genai.configure(api_key=GOOGLE_API_KEY)
    gemini_model = genai.GenerativeModel('gemini-2.5-flash')
except Exception as e:
    st.error(f"Gemini API Hatası: {e}")

try:
    groq_client = Groq(api_key=GROQ_API_KEY)
except Exception as e:
    st.error(f"Groq API Hatası: {e}")

# --- 2. YARDIMCI FONKSİYONLAR ---

def tr_duzelt(metin):
    """Sadece görüntüleme için basit karakter düzeltme."""
    dic = {'ğ':'g', 'Ğ':'G', 'ş':'s', 'Ş':'S', 'ı':'i', 'İ':'I', 'ç':'c', 'Ç':'C', 'ü':'u', 'Ü':'U', 'ö':'o', 'Ö':'O'}
    for k, v in dic.items():
        metin = metin.replace(k, v)
    return metin

# 3. WORD FONKSİYONU (SINAV ASİSTANI İÇİN)
def create_exam_word(sorular_kismi, cevaplar_kismi):
    doc = Document()
    doc.add_heading('SINAV KAĞIDI', 0)
    doc.add_paragraph(sorular_kismi)
    doc.add_page_break()
    doc.add_heading('CEVAP ANAHTARI', 1)
    doc.add_paragraph(cevaplar_kismi)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()

# 4. WORD FONKSİYONU (TOPLANTI ASİSTANI İÇİN)
def create_meeting_word(tutanak_metni, transkript_metni):
    doc = Document()
    doc.add_heading('TOPLANTI TUTANAĞI RAPORU', 0)
    doc.add_heading('1. YAPAY ZEKA ÖZETİ', 1)
    doc.add_paragraph(tutanak_metni)
    doc.add_page_break()
    doc.add_heading('2. ORİJİNAL KONUŞMA DÖKÜMÜ (TRANSKRİPT)', 1)
    doc.add_paragraph(transkript_metni)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


# 5. CLEAR STATE
def meeting_clear_state():
    st.session_state.meeting_tutanak = None
    st.session_state.meeting_transkript = None


# --- 6. ANA SAYFA VE TABLAR (LOGOSUZ YENİ TASARIM) ---
st.set_page_config(
    page_title="Maarif Suite",
    page_icon="🎓",
    layout="wide" 
)

# LOGO KALDIRILDI, BAŞLIKLAR BÜYÜTÜLDÜ VE ORTALANDI

# Büyük Başlık (H1, font-size: 3.5em)
st.markdown(
    "<h1 style='text-align: center; color: #1E3A8A; font-size: 3.5em;'>MAARİF SUITE</h1>", 
    unsafe_allow_html=True
)

# Alt Başlık (P, font-size: 1.3em)
st.markdown(
    "<p style='text-align: center; color: gray; font-size: 1.3em;'>Eğitim Teknolojilerinde İki Güç Bir Arada</p>", 
    unsafe_allow_html=True
)
st.write("---") 

tab_exam, tab_meeting, tab_about = st.tabs(["🎓 SINAV ASİSTANI (Gemini)", "🎙️ TOPLANTI ASİSTANI (Groq)", "ℹ️ HAKKINDA"])

# ----------------------------------------------------------------------
#                         TAB 3: HAKKINDA
# ----------------------------------------------------------------------

with tab_about:
    st.header("Vizyonumuz ve Hakkımda")
    st.subheader("👨‍💻 Geliştirici: Nejdet TUT")
    
    st.markdown(f"""
    Merhaba, ben **Nejdet TUT**. Uzman bir **Bilişim Teknolojileri Öğretmeni** ve **EdTech Geliştiricisiyim**. Grafik tasarım kökenli bir teknoloji eğitimcisi olarak, **12 yılı aşkın öğretmenlik** tecrübemi Yapay Zeka ve Veri Bilimi ile birleştiriyorum.

    **Eğitim Bilgisi:** Trakya Üniversitesi'nden Bilgisayar ve Öğretim Teknolojileri Öğretmenliği bölümünden mezun oldum.
    """)
    
    st.subheader("💡 Proje Amacı: Öğretmen Verimliliğini Artırmak")
    st.markdown("""
    **Maarif Suite**, öğretmenlerin üzerindeki idari ve hazırlık yükünü hafifletmek için tasarlanmıştır. Uygulamanın temel hedefleri şunlardır:
    * **Sınav Otomasyonu:** Gemini API gücüyle müfredata uyumlu sınav sorularını otomatik olarak oluşturarak hazırlık süresini **%90 oranında** azaltmak.
    * **Zaman Yönetimi:** Toplantı ve ders dökümlerini anında analiz ederek profesyonel tutanaklar hazırlamak (Groq/Whisper ile).
    """)
    
    st.subheader("📞 İletişim Bilgileri")
    st.markdown(f"""
    * **E-posta:** nejdettut@gmail.com
    * **Telefon:** +90 507 795 79 36
    * **LinkedIn:** [linkedin.com/in/nejdettut](https://www.linkedin.com/in/nejdettut)
    """)

# ----------------------------------------------------------------------
#                         TAB 1: SINAV ASİSTANI
# ----------------------------------------------------------------------

with tab_exam:
    st.markdown("### ✨ Yapay Zeka Destekli Sınav Kurgulama (Word İndirme)")
    
    with st.expander("⚙️ Sınav Ayarlarını Yapılandır (Tıkla)", expanded=False):
        c1, c2, c3, c4 = st.columns(4)
        with c1:
            olcum_turu = st.selectbox("Ölçme Türü:", ("Çoktan Seçmeli", "Doğru/Yanlış", "Klasik", "Boşluk Doldurma", "Eşleştirme"), key="olcum_turu")
        with c2:
            seviye = st.selectbox("Sınıf Seviyesi:", ("İlkokul (1-4)", "Ortaokul (5-8)", "Lise (9-12)", "Üniversite Hazırlık"), key="exam_level")
        with c3:
            zorluk = st.slider("Zorluk:", 1, 5, 3, key="exam_diff")
        with c4:
            soru_sayisi = st.number_input("Soru Sayısı:", 1, 20, 5, key="exam_count")

    konu = st.text_input("", placeholder="Hangi konuda sınav hazırlamak istersin?", key="exam_topic")
    generate_btn = st.button("✨ Sınavı Oluştur", key="exam_gen", type="primary", use_container_width=True)

    if generate_btn:
        if not konu: st.warning("Lütfen bir konu yazın.")
        else:
            with st.spinner('Yapay Zeka soruları kurguluyor...'):
                try:
                    prompt = f"""
                    Sen MEB müfredatına hakim uzman bir öğretmensin.
                    Konu: {konu}, Seviye: {seviye}, Zorluk: {zorluk}/5, Soru Sayısı: {soru_sayisi}.
                    Sınav Türü: {olcum_turu}.

                    GÖREV: Soruları istenen formatta hazırlarken, öğrencilerin seviyesine uygun ve MEB müfredatına hakim ol.
                    EN SONA, sorular bittikten sonra tam olarak şu ayırıcıyı koy: "---CEVAP_ANAHTARI_BOLUMU---"
                    Bu ayırıcıdan sonra cevap anahtarını yaz.
                    """
                    response = gemini_model.generate_content(prompt)
                    full_text = response.text
                    
                    if "---CEVAP_ANAHTARI_BOLUMU---" in full_text:
                        parts = full_text.split("---CEVAP_ANAHTARI_BOLUMU---")
                        sorular_kismi = parts[0].strip()
                        cevaplar_kismi = parts[1].strip()
                    else:
                        sorular_kismi = full_text
                        cevaplar_kismi = "Cevap anahtarı ayrıştırılamadı."

                    st.success("Sınav Hazır!")
                    st.write(sorular_kismi)
                    with st.expander("Cevap Anahtarını Gör"): st.write(cevaplar_kismi)
                    
                    word_data = create_exam_word(sorular_kismi, cevaplar_kismi)

                    st.download_button(
                        label="📑 Cevap Anahtarlı İndir (Word)",
                        data=word_data,
                        file_name=f"{konu}_sinav.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                        type="secondary"
                    )

                except Exception as e:
                    st.error(f"Sınav Oluşturma Hatası: {e}")

# ----------------------------------------------------------------------
#                      TAB 2: TOPLANTI ASİSTANI
# ----------------------------------------------------------------------

with tab_meeting:
    st.markdown("### 🎙️ Sesli Toplantı Tutanak Motoru")
    
    if 'meeting_tutanak' not in st.session_state: st.session_state.meeting_tutanak = None
    if 'meeting_transkript' not in st.session_state: st.session_state.meeting_transkript = None
    
    col_upload, col_record = st.columns([1, 1])
    with col_upload:
        uploaded_file = st.file_uploader("Ses Dosyası Yükle (mp3, wav)", type=['mp3', 'wav', 'm4a'], key="meeting_upload")
    with col_record:
        audio_recording = st.audio_input("Canlı Kayıt Başlat", key="meeting_record")

    ses_verisi = uploaded_file if uploaded_file else audio_recording
    
    analiz_yapildi = st.session_state.meeting_tutanak is not None

    # --- İŞLEM KISMI ---
    if ses_verisi:
        st.write("---")
        
        col_start, col_reset = st.columns(2)

        with col_start:
            # Analizi Başlat Butonu (Tasarım: Sonuç varsa devre dışı kalır)
            if st.button("📝 Analizi Başlat", key="meeting_start", type="primary", use_container_width=True, disabled=analiz_yapildi):
                with st.spinner("⚡ Groq/Whisper motoru dinliyor ve Llama 3 analiz ediyor..."):
                    try:
                        # [ANALİZ KODU BAŞLANGIÇ]
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as tmp_file:
                            tmp_file.write(ses_verisi.getvalue())
                            tmp_file_path = tmp_file.name

                        with open(tmp_file_path, "rb") as file:
                            transcription_result = groq_client.audio.transcriptions.create(
                                file=(tmp_file_path, file.read()),
                                model="whisper-large-v3",
                                response_format="text"
                            )
                        st.session_state.meeting_transkript = transcription_result
                        
                        prompt = f"""
                        Aşağıdaki metin bir toplantı dökümüdür. Bunu profesyonel bir tutanak haline getir.
                        METİN: {st.session_state.meeting_transkript}
                        İSTENEN RAPOR FORMATI: 1. 📝 ÖZET 2. ✅ ALINAN KARARLAR 3. 📌 GÖREV DAĞILIMI
                        """
                        completion = groq_client.chat.completions.create(
                            model="llama-3.3-70b-versatile",
                            messages=[{"role": "system", "content": "Sen profesyonel bir okul asistanısın. Türkçe cevap ver."}, {"role": "user", "content": prompt}],
                        )
                        st.session_state.meeting_tutanak = completion.choices[0].message.content
                        os.remove(tmp_file_path)
                        st.rerun() # Sayfayı yenileyip sonucu göster

                    except Exception as e:
                        st.error(f"Analiz Hatası: {e}")

        with col_reset:
            # Analizi Sıfırla Butonu (Aynı hizada, aynı stil)
            st.button("🔄 Analizi Sıfırla / Yeni Ses", on_click=meeting_clear_state, key="meeting_reset_col", type="secondary", use_container_width=True)

    # --- SONUÇLARI GÖSTER VE KAYDET BUTONU ---
    if st.session_state.meeting_tutanak is not None:
        st.write("---")
        st.success("Analiz Başarılı! Raporu inceleyip aşağıdan indirebilirsiniz.")

        with st.expander("📄 Konuşma Dökümünü Gör (Transkript)", expanded=False):
            st.write(st.session_state.meeting_transkript)
        
        st.markdown("### 📋 Oluşturulan Tutanak")
        st.markdown(st.session_state.meeting_tutanak)
        
        st.write("---")

        word_data = create_meeting_word(st.session_state.meeting_tutanak, st.session_state.meeting_transkript)
        
        st.download_button(
            label="Analizi Kaydet (Word)",
            data=word_data,
            file_name="toplanti_tutanagi.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            type="primary"
        )
